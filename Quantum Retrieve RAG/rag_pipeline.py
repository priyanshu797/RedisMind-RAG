import os
import hashlib
import logging
import tempfile
import zipfile
import json
import csv
import re
import time
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from collections import OrderedDict
import pickle
from dotenv import load_dotenv
import nltk
import docx
import whisper
import pdfplumber
import pytesseract
import cv2
import numpy as np

from nltk.tokenize import sent_tokenize
from llama_index.core import VectorStoreIndex, Document, StorageContext, load_index_from_storage, Settings
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.groq import Groq
from llama_index.llms.ollama import Ollama
from sentence_transformers import CrossEncoder

import redis
from redis.commands.search.field import VectorField, TextField, NumericField
from redis.commands.search.indexDefinition import IndexDefinition, IndexType
from redis.commands.search.query import Query

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(BASE_DIR, ".env"))

nltk.download("punkt", quiet=True)
nltk.download("punkt_tab", quiet=True)

#Configuration
PERSIST_DIR = "./storage"
CACHE_DIR = "./cache"
HASH_CACHE_FILE = "text_hash.txt"
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
GROQ_MODEL = "llama-3.3-70b-versatile"
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
OLLAMA_MODEL = "gemma2:2b"
CHUNK_SIZE = 512
CHUNK_OVERLAP = 50

# Retrieval parameters
K_INIT = 20
K_RERANK = 10
K_FINAL = 5
SIMILARITY_THRESHOLD = 0.7
CACHE_SIZE = 1000
CACHE_SIMILARITY_THRESHOLD = 0.95

#Redis configuration
REDIS_HOST = os.getenv("REDIS_HOST", "localhost")
REDIS_PORT = int(os.getenv("REDIS_PORT", 6379))
REDIS_PASSWORD = os.getenv("REDIS_PASSWORD", None)
REDIS_DB = int(os.getenv("REDIS_DB", 0))
REDIS_CACHE_INDEX = "rag_semantic_cache"
REDIS_CACHE_PREFIX = "rag:cache:" 
REDIS_CACHE_TTL = int(os.getenv("REDIS_CACHE_TTL", 86400))
EMBEDDING_DIM = 384      

logging.basicConfig(level=logging.WARNING, format="%(asctime)s - %(levelname)s - %(message)s")

@dataclass
class CacheEntry:
    """Minimal payload stored per cached query."""
    query: str
    response: str
    sources: List[str]
    timestamp: float


class RedisSemanticCache:

    def __init__(
        self,
        host: str = REDIS_HOST,
        port: int = REDIS_PORT,
        password: Optional[str] = REDIS_PASSWORD,
        db: int = REDIS_DB,
        max_l1_size: int = 256,
        similarity_threshold: float = CACHE_SIMILARITY_THRESHOLD,
        ttl: int = REDIS_CACHE_TTL,
    ):
        self.similarity_threshold = similarity_threshold
        self.max_l1_size = max_l1_size
        self.ttl = ttl
        self._l1: OrderedDict[str, Tuple[np.ndarray, CacheEntry]] = OrderedDict()
        self._redis: Optional[redis.Redis] = None
        self._redis_ok = False
        self._connect(host, port, password, db)

        if self._redis_ok:
            self._ensure_index()


    def _connect(self, host, port, password, db):
        try:
            self._redis = redis.Redis(
                host=host,
                port=port,
                password=password,
                db=db,
                decode_responses=False,
                socket_connect_timeout=2,
                socket_timeout=2,
            )
            self._redis.ping()
            self._redis_ok = True
            logging.info(f"✓ Redis connected at {host}:{port}")
        except Exception as exc:
            logging.warning(
                f"Redis unavailable ({exc}). "
                "Make sure redis-stack-server is installed and running. "
                "Falling back to L1 (in-process) cache only."
            )
            self._redis_ok = False

    def _ensure_index(self):
        """Create the HNSW vector index if it does not already exist."""
        try:
            self._redis.ft(REDIS_CACHE_INDEX).info()
            logging.info(f"✓ RediSearch index '{REDIS_CACHE_INDEX}' already exists.")
        except Exception:
            schema = [
                TextField("query"),
                TextField("response"),
                TextField("sources"),
                NumericField("timestamp"),
                VectorField(
                    "embedding",
                    "HNSW",
                    {
                        "TYPE": "FLOAT32",
                        "DIM": EMBEDDING_DIM,
                        "DISTANCE_METRIC": "COSINE",
                        "M": 16,
                        "EF_CONSTRUCTION": 200,
                    },
                ),
            ]
            self._redis.ft(REDIS_CACHE_INDEX).create_index(
                schema,
                definition=IndexDefinition(
                    prefix=[REDIS_CACHE_PREFIX],
                    index_type=IndexType.HASH,
                ),
            )
            logging.info(f"✓ Created RediSearch HNSW index '{REDIS_CACHE_INDEX}'.")

    @staticmethod
    def _md5(text: str) -> str:
        return hashlib.md5(text.encode()).hexdigest()

    @staticmethod
    def _to_bytes(vec: np.ndarray) -> bytes:
        return vec.astype(np.float32).tobytes()

    @staticmethod
    def _from_bytes(raw: bytes) -> np.ndarray:
        return np.frombuffer(raw, dtype=np.float32)

    @staticmethod
    def _cosine(a: np.ndarray, b: np.ndarray) -> float:
        denom = np.linalg.norm(a) * np.linalg.norm(b)
        return float(np.dot(a, b) / denom) if denom else 0.0

    def _l1_touch(self, key: str):
        """Move key to end (most-recently-used) in L1."""
        if key in self._l1:
            self._l1.move_to_end(key)

    def _l1_evict(self):
        while len(self._l1) >= self.max_l1_size:
            self._l1.popitem(last=False)  # drop LRU

    def get(
        self,
        query: str,
        query_embedding: Optional[np.ndarray] = None,
    ) -> Optional[Tuple[str, List[str]]]:
        key = self._md5(query)

        # 1. L1 exact
        if key in self._l1:
            self._l1_touch(key)
            _, entry = self._l1[key]
            logging.info("Cache HIT  L1-exact")
            return entry.response, entry.sources

        # 2. L1 semantic
        if query_embedding is not None:
            for k, (emb, entry) in self._l1.items():
                if self._cosine(query_embedding, emb) >= self.similarity_threshold:
                    self._l1_touch(k)
                    logging.info("Cache HIT  L1-semantic")
                    return entry.response, entry.sources

        if not self._redis_ok:
            return None

        # 3. L2 HNSW vector search
        if query_embedding is not None:
            try:
                q = (
                    Query("*=>[KNN 5 @embedding $vec AS score]")
                    .sort_by("score")
                    .return_fields("query", "response", "sources", "timestamp", "embedding", "score")
                    .dialect(2)
                )
                results = self._redis.ft(REDIS_CACHE_INDEX).search(
                    q, query_params={"vec": self._to_bytes(query_embedding)}
                )
                for doc in results.docs:
                    distance = float(getattr(doc, "score", 2.0))
                    sim = 1.0 - distance
                    if sim >= self.similarity_threshold:
                        response = doc.response
                        sources = json.loads(doc.sources)
                        emb_cached = self._from_bytes(doc.embedding)
                        entry = CacheEntry(
                            query=doc.query,
                            response=response,
                            sources=sources,
                            timestamp=float(doc.timestamp),
                        )
                        # Promote to L1
                        self._l1_evict()
                        self._l1[key] = (emb_cached, entry)
                        logging.info(f"Cache HIT  L2-HNSW (sim={sim:.3f})")
                        return response, sources
            except Exception as exc:
                logging.warning(f"Redis HNSW search failed: {exc}")
        try:
            redis_key = f"{REDIS_CACHE_PREFIX}{key}"
            raw = self._redis.hgetall(redis_key)
            if raw:
                response = raw[b"response"].decode()
                sources = json.loads(raw[b"sources"].decode())
                emb_cached = self._from_bytes(raw[b"embedding"])
                entry = CacheEntry(
                    query=raw[b"query"].decode(),
                    response=response,
                    sources=sources,
                    timestamp=float(raw[b"timestamp"]),
                )
                self._l1_evict()
                self._l1[key] = (emb_cached, entry)
                logging.info("Cache HIT  L2-exact")
                return response, sources
        except Exception as exc:
            logging.warning(f"Redis exact lookup failed: {exc}")

        return None

    def set(
        self,
        query: str,
        query_embedding: np.ndarray,
        response: str,
        sources: List[str],
    ):
        """
        Store a new cache entry in both L1 and Redis (L2).
        """
        key = self._md5(query)
        entry = CacheEntry(
            query=query,
            response=response,
            sources=sources,
            timestamp=time.time(),
        )

        # L1
        self._l1_evict()
        self._l1[key] = (query_embedding, entry)

        # L2 – Redis Hash (indexed by RediSearch)
        if not self._redis_ok:
            return
        try:
            redis_key = f"{REDIS_CACHE_PREFIX}{key}"
            mapping = {
                "query": query,
                "response": response,
                "sources": json.dumps(sources),
                "timestamp": str(entry.timestamp),
                "embedding": self._to_bytes(query_embedding),
            }
            pipe = self._redis.pipeline(transaction=False)
            pipe.hset(redis_key, mapping=mapping)
            if self.ttl > 0:
                pipe.expire(redis_key, self.ttl)
            pipe.execute()
            logging.info(f"Cache SET  key={key[:8]}…  TTL={self.ttl}s")
        except Exception as exc:
            logging.warning(f"Redis SET failed: {exc}")

    def invalidate(self, query: str):
        """Remove a specific query from both L1 and L2."""
        key = self._md5(query)
        self._l1.pop(key, None)
        if self._redis_ok:
            try:
                self._redis.delete(f"{REDIS_CACHE_PREFIX}{key}")
            except Exception as exc:
                logging.warning(f"Redis DELETE failed: {exc}")

    def flush(self):
        """Clear all cache entries (L1 + L2 namespace)."""
        self._l1.clear()
        if self._redis_ok:
            try:
                keys = self._redis.keys(f"{REDIS_CACHE_PREFIX}*")
                if keys:
                    self._redis.delete(*keys)
                logging.info("Cache flushed (L1 + L2).")
            except Exception as exc:
                logging.warning(f"Redis FLUSH failed: {exc}")

    @property
    def l1_size(self) -> int:
        return len(self._l1)

    def stats(self) -> Dict:
        info = {"l1_entries": self.l1_size, "redis_ok": self._redis_ok}
        if self._redis_ok:
            try:
                idx_info = self._redis.ft(REDIS_CACHE_INDEX).info()
                info["l2_entries"] = idx_info.get("num_docs", "?")
            except Exception:
                info["l2_entries"] = "?"
        return info

cache = RedisSemanticCache()

class QueryClassifier:
    FACTUAL_KEYWORDS = ["what", "who", "when", "where", "define", "explain", "list", "show", "tell"]
    ANALYTICAL_KEYWORDS = ["compare", "analyze", "why", "how can", "how does", "evaluate", "difference", "relationship"]
    CONVERSATIONAL_KEYWORDS = ["hi", "hello", "thanks", "thank you", "bye", "hey"]

    @staticmethod
    def classify(query: str) -> Dict[str, any]:
        query_lower = query.lower().strip()
        word_count = len(query.split())

        is_pure_conversational = (
            word_count <= 3
            and any(query_lower.startswith(kw) for kw in QueryClassifier.CONVERSATIONAL_KEYWORDS)
        )

        intent = "factual"
        if is_pure_conversational:
            intent = "conversational"
        elif any(kw in query_lower for kw in QueryClassifier.ANALYTICAL_KEYWORDS):
            intent = "analytical"
        elif any(kw in query_lower for kw in QueryClassifier.FACTUAL_KEYWORDS):
            intent = "factual"

        complexity = min(1.0, (word_count / 20) + (0.3 if intent == "analytical" else 0))

        return {"intent": intent, "complexity": complexity, "word_count": word_count}

def initialize_settings(use_groq=True):
    try:
        Settings.embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)
        Settings.chunk_size = CHUNK_SIZE
        Settings.chunk_overlap = CHUNK_OVERLAP

        if use_groq:
            try:
                if not GROQ_API_KEY or GROQ_API_KEY == "your_groq_api_key_here":
                    raise ValueError("Missing GROQ_API_KEY")
                Settings.llm = Groq(
                    model=GROQ_MODEL,
                    api_key=GROQ_API_KEY,
                    temperature=0.3,
                    max_tokens=600,
                )
                logging.info(f"Using Groq LLM: {GROQ_MODEL}")
                return "groq"
            except Exception as e:
                logging.warning(f"Groq failed: {e}. Falling back to Ollama…")
                raise
    except Exception:
        Settings.llm = Ollama(model=OLLAMA_MODEL, temperature=0.3)
        logging.info(f"Using Ollama LLM: {OLLAMA_MODEL}")
        return "ollama"

def extract_text_from_pdf(path):
    try:
        with pdfplumber.open(path) as pdf:
            return "\n".join([p.extract_text() or "" for p in pdf.pages])
    except Exception as e:
        logging.error(f"PDF extraction failed: {e}")
        return ""

def extract_text_from_image(path):
    try:
        img = cv2.imread(path)
        return pytesseract.image_to_string(img)
    except Exception as e:
        logging.error(f"Image OCR failed: {e}")
        return ""

def extract_text_from_docx(path):
    try:
        return "\n".join([p.text for p in docx.Document(path).paragraphs])
    except Exception as e:
        logging.error(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_txt(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        logging.error(f"TXT extraction failed: {e}")
        return ""

def extract_text_from_csv(path):
    try:
        with open(path, newline='', encoding='utf-8') as csvfile:
            return "\n".join([" | ".join(row) for row in csv.reader(csvfile)])
    except Exception as e:
        logging.error(f"CSV extraction failed: {e}")
        return ""

def extract_text_from_json(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.dumps(json.load(f), indent=2)
    except Exception as e:
        logging.error(f"JSON extraction failed: {e}")
        return ""

def extract_text_from_audio(path):
    try:
        model = whisper.load_model("base")
        return model.transcribe(path)["text"]
    except Exception as e:
        logging.error(f"Audio transcription failed: {e}")
        return ""

def extract_text_from_zip(path):
    text = ""
    try:
        with zipfile.ZipFile(path, 'r') as zip_ref:
            with tempfile.TemporaryDirectory() as temp_dir:
                zip_ref.extractall(temp_dir)
                for root, _, files in os.walk(temp_dir):
                    for name in files:
                        full_path = os.path.join(root, name)
                        text += extract_text(full_path) + "\n"
    except Exception as e:
        logging.error(f"ZIP extraction failed: {e}")
    return text

def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    extractors = {
        ".pdf": extract_text_from_pdf,
        ".jpg": extract_text_from_image,
        ".jpeg": extract_text_from_image,
        ".png": extract_text_from_image,
        ".docx": extract_text_from_docx,
        ".txt": extract_text_from_txt,
        ".csv": extract_text_from_csv,
        ".json": extract_text_from_json,
        ".mp3": extract_text_from_audio,
        ".wav": extract_text_from_audio,
        ".mp4": extract_text_from_audio,
        ".zip": extract_text_from_zip,
    }
    extractor = extractors.get(ext)
    return extractor(path) if extractor else ""

def clean_text(text):
    return re.sub(r'\s+', ' ', text).replace("•", "-").replace("–", "-").strip()

def compute_hash(text):
    return hashlib.md5(text.encode("utf-8")).hexdigest()

def has_file_changed(text):
    if os.path.exists(HASH_CACHE_FILE):
        with open(HASH_CACHE_FILE, "r") as f:
            return compute_hash(text) != f.read().strip()
    return True

def cache_hash(text):
    with open(HASH_CACHE_FILE, "w") as f:
        f.write(compute_hash(text))

def compress_context(chunks: List[str], max_tokens: int = 1500) -> str:
    compressed = []
    token_count = 0

    for chunk in chunks:
        chunk_tokens = len(chunk.split())
        if token_count + chunk_tokens > max_tokens:
            remaining = max_tokens - token_count
            compressed.append(" ".join(chunk.split()[:remaining]))
            break
        compressed.append(chunk)
        token_count += chunk_tokens

    unique_sentences = []
    seen = set()
    for chunk in compressed:
        for sent in sent_tokenize(chunk):
            sent_clean = sent.strip().lower()
            if sent_clean not in seen and len(sent_clean) > 20:
                seen.add(sent_clean)
                unique_sentences.append(sent)

    return " ".join(unique_sentences)

def create_index(file_paths):
    documents = []
    for path in file_paths:
        text = clean_text(extract_text(path))
        if text:
            doc = Document(text=text, metadata={"source": os.path.basename(path)})
            documents.append(doc)

    if not documents:
        logging.warning("No valid text extracted from files.")
        return None

    combined_text = "\n".join([doc.text for doc in documents])

    if os.path.exists(PERSIST_DIR) and not has_file_changed(combined_text):
        logging.info("Loading existing index…")
        storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
        index = load_index_from_storage(storage_context)
    else:
        logging.info("Creating new index…")
        index = VectorStoreIndex.from_documents(documents, show_progress=True)
        index.storage_context.persist(persist_dir=PERSIST_DIR)
        cache_hash(combined_text)

    return index

def format_response(response: str, sources: List[str]) -> str:
    response = response.strip()
    response = re.sub(r'^(Answer|Response|Here\'s|Based on):\s*', '', response, flags=re.IGNORECASE)
    response = re.sub(r'\*\*([^*]+)\*\*', r'\1', response)
    response = re.sub(r'__([^_]+)__', r'\1', response)
    response = re.sub(r'\n{3,}', '\n\n', response)

    if sources:
        unique_sources = list(dict.fromkeys(sources))
        response += f" [Source: {', '.join(unique_sources)}]"

    return response

def query_index_optimized(
    index, query: str, llm_type: str
) -> Tuple[str, List[str], Dict[str, float]]:
    timings: Dict[str, float] = {}
    start_total = time.time()

    # Query Classification
    t0 = time.time()
    classification = QueryClassifier.classify(query)
    timings["classification"] = time.time() - t0

    t0 = time.time()
    query_embedding = np.array(Settings.embed_model.get_text_embedding(query), dtype=np.float32)
    cached = cache.get(query, query_embedding)
    timings["cache_lookup"] = time.time() - t0

    if cached:
        response, sources = cached
        timings["total"] = time.time() - start_total
        return response, sources, timings

    # Conversational
    if classification["intent"] == "conversational":
        response = "Hello! Ask me anything about your documents."
        cache.set(query, query_embedding, response, [])
        timings["total"] = time.time() - start_total
        return response, [], timings

    t0 = time.time()
    k_retrieve = K_INIT if classification["complexity"] > 0.5 else max(K_FINAL, 8)
    retriever = index.as_retriever(similarity_top_k=k_retrieve)
    retrieved_nodes = retriever.retrieve(query)
    timings["retrieval"] = time.time() - t0

    if not retrieved_nodes:
        response = "I couldn't find relevant information in the documents to answer that."
        timings["total"] = time.time() - start_total
        return response, [], timings

    # Reranking
    t0 = time.time()
    top_score = retrieved_nodes[0].score if hasattr(retrieved_nodes[0], "score") else 0

    if top_score < SIMILARITY_THRESHOLD and len(retrieved_nodes) > K_FINAL:
        try:
            reranker = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")
            pairs = [[query, node.text] for node in retrieved_nodes[:K_RERANK]]
            scores = reranker.predict(pairs)
            scored_nodes = sorted(
                zip(retrieved_nodes[:K_RERANK], scores), key=lambda x: x[1], reverse=True
            )
            final_nodes = [node for node, _ in scored_nodes[:K_FINAL]]
        except Exception:
            final_nodes = retrieved_nodes[:K_FINAL]
    else:
        final_nodes = retrieved_nodes[:K_FINAL]

    timings["reranking"] = time.time() - t0

    t0 = time.time()
    contexts = [node.text for node in final_nodes]
    compressed_context = compress_context(contexts, max_tokens=1200)
    timings["compression"] = time.time() - t0

    t0 = time.time()
    prompt = (
        "You are an AI assistant. Use only the given context to answer.\n"
        "If answer is not in context, reply: 'Not available in the documentation.'\n\n"
        f"Context:\n{compressed_context}\n\n"
        f"Question: {query}\n"
        "Answer:"
    )
    response_obj = Settings.llm.complete(prompt)
    response = response_obj.text.strip()

    generic_phrases = [
        "in general", "typically", "usually", "commonly", "it is known that",
        "experts suggest", "research shows", "studies indicate",
    ]
    if any(p in response.lower() for p in generic_phrases) and len(response.split()) < 30:
        response = "This information is not available in the provided documents."

    timings["generation"] = time.time() - t0

    # Source extraction + cache write  ← RedisSemanticCache.set()
    sources = [node.metadata.get("source", "Unknown") for node in final_nodes[:3]]
    cache.set(query, query_embedding, response, sources)

    timings["total"] = time.time() - start_total
    return response, sources, timings

def main():
    try:
        llm_type = initialize_settings(use_groq=True)
    except Exception:
        llm_type = initialize_settings(use_groq=False)

    # Show cache stats on startup
    print(f"Cache stats: {cache.stats()}")

    file_input = input("Enter file paths (comma-separated): ").strip()
    file_paths = [f.strip() for f in file_input.split(",") if os.path.exists(f.strip())]

    if not file_paths:
        print("No valid file paths provided.")
        return

    print(f"\nProcessing {len(file_paths)} file(s)…")
    index = create_index(file_paths)

    if index is None:
        print("Failed to create index.")
        return

    while True:
        question = input("Q: ").strip()
        if question.lower() in ["exit", "quit", "q"]:
            print("\nGoodbye!")
            break
        if not question:
            continue
        response, sources, timings = query_index_optimized(index, question, llm_type)
        formatted_response = format_response(response, sources)
        print(f"\n{formatted_response}\n")
        print(f"({timings['total']*1000:.0f}ms)\n")

if __name__ == "__main__":
    main()